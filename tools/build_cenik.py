from pathlib import Path
from copy import deepcopy

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
PDF_DIR = ROOT / "assets" / "pdf"
IMG_DIR = ROOT / "assets" / "img"
LOGO = IMG_DIR / "logo.png"
WATERMARK = PDF_DIR / "cenik-logo-watermark.png"
DOCX_OUT = PDF_DIR / "cenik.docx"
PDF_OUT = PDF_DIR / "cenik.pdf"

RED = "C91722"
DARK = "111318"
GRAY = "6F747C"
LIGHT = "F6F7F9"
BORDER = "363A40"
PDF_FONT = "ArialAuto"
PDF_FONT_BOLD = "ArialAutoBold"


ITEMS_BASIC = [
    ("Mechanické práce", "750 Kč/h"),
    ("Karosářské práce", "750 Kč/h"),
    ("Elektrikářské práce", "850 Kč/h"),
]

ITEMS_DIAGNOSTICS = [
    ("Seriová diagnostika", "Načtení závad a jejich výpis.", "500 Kč"),
    (
        "Paralelní diagnostika",
        "Měření, bloky načtených hodnot, test jízdy s diagnostikou, vyhledávání závady.",
        "500 Kč + 750 Kč/h",
    ),
]

ITEMS_AC = [
    ("R134A", ""),
    ("R1234yf", ""),
    ("Detekce úniku chladiva z okruhu klimatizace", "1500 Kč"),
]


def make_watermark() -> None:
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    logo = Image.open(LOGO).convert("RGBA")
    logo.thumbnail((1450, 640), Image.Resampling.LANCZOS)
    alpha = logo.getchannel("A").point(lambda p: int(p * 0.12))
    logo.putalpha(alpha)
    canvas = Image.new("RGBA", (1700, 850), (255, 255, 255, 0))
    x = (canvas.width - logo.width) // 2
    y = (canvas.height - logo.height) // 2
    canvas.alpha_composite(logo, (x, y))
    canvas.save(WATERMARK)


def register_pdf_fonts() -> None:
    font_dir = Path("C:/Windows/Fonts")
    pdfmetrics.registerFont(TTFont(PDF_FONT, str(font_dir / "arial.ttf")))
    pdfmetrics.registerFont(TTFont(PDF_FONT_BOLD, str(font_dir / "arialbd.ttf")))


def set_font(run, size=None, color=None, bold=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=120, start=160, bottom=120, end=160):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color=RED, size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_logo_to_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(LOGO), width=Cm(4.4))


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=14, color=DARK, bold=True)
    add_bottom_border(p, RED, "8")


def fill_cell(cell, text, bold=False, color=DARK, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.35)
    section.footer_distance = Cm(0.8)
    add_logo_to_header(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("Auto Sosnovec s.r.o.")
    set_font(run, size=11, color=GRAY, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("Ceník služeb")
    set_font(run, size=28, color=DARK, bold=True)
    add_bottom_border(title, RED, "18")

    add_heading(doc, "Mechanické, karosářské a elektrikářské práce")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(11.2)
    table.columns[1].width = Cm(5.4)
    set_cell_margins(table)
    headers = ("Služba", "Cena")
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, DARK)
        set_cell_border(cell, DARK)
        fill_cell(cell, text, bold=True, color="FFFFFF", size=10.5)
    for service, price in ITEMS_BASIC:
        row = table.add_row().cells
        for cell in row:
            set_cell_border(cell, "B8BDC5", "4")
        fill_cell(row[0], service, bold=True)
        fill_cell(row[1], price, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    add_heading(doc, "Diagnostika")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(8.8)
    table.columns[2].width = Cm(3.3)
    set_cell_margins(table)
    for cell, text in zip(table.rows[0].cells, ("Typ", "Rozsah práce", "Cena")):
        set_cell_shading(cell, DARK)
        set_cell_border(cell, DARK)
        fill_cell(cell, text, bold=True, color="FFFFFF", size=10.5)
    for service, description, price in ITEMS_DIAGNOSTICS:
        row = table.add_row().cells
        for cell in row:
            set_cell_border(cell, "B8BDC5", "4")
        fill_cell(row[0], service, bold=True)
        fill_cell(row[1], description)
        fill_cell(row[2], price, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    add_heading(doc, "Plnění klimatizací")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(11.2)
    table.columns[1].width = Cm(5.4)
    set_cell_margins(table)
    for cell, text in zip(table.rows[0].cells, ("Služba", "Cena")):
        set_cell_shading(cell, DARK)
        set_cell_border(cell, DARK)
        fill_cell(cell, text, bold=True, color="FFFFFF", size=10.5)
    for service, price in ITEMS_AC:
        row = table.add_row().cells
        for cell in row:
            set_cell_border(cell, "B8BDC5", "4")
        fill_cell(row[0], service, bold=True)
        fill_cell(row[1], price, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(16)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Ceny položek R134A a R1234yf jsou ponechány k ručnímu doplnění ve Word kopii.")
    set_font(run, size=9.5, color=GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Auto Sosnovec s.r.o. | Přístavní 27, Praha 7 | autososnovec@seznam.cz")
    set_font(run, size=8.5, color=GRAY)

    doc.save(DOCX_OUT)


def make_pdf_table(data, col_widths, header_cols=2):
    table = Table(data, colWidths=col_widths, hAlign="CENTER")
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111318")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), PDF_FONT),
        ("FONTNAME", (0, 0), (-1, 0), PDF_FONT_BOLD),
        ("FONTSIZE", (0, 0), (-1, -1), 9.4),
        ("FONTNAME", (0, 1), (0, -1), PDF_FONT_BOLD),
        ("FONTNAME", (-1, 1), (-1, -1), PDF_FONT_BOLD),
        ("ALIGN", (-1, 1), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.55, colors.HexColor("#B8BDC5")),
        ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#363A40")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]
    table.setStyle(TableStyle(style))
    return table


def build_pdf():
    register_pdf_fonts()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName=PDF_FONT,
        fontSize=9.4,
        leading=12,
        textColor=colors.HexColor("#111318"),
    )
    section_title = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontName=PDF_FONT_BOLD,
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#111318"),
        spaceBefore=14,
        spaceAfter=8,
        borderColor=colors.HexColor("#C91722"),
        borderWidth=0,
        borderPadding=0,
    )

    def draw_background(canvas, _doc):
        canvas.saveState()
        page_w, page_h = A4
        canvas.setFillColor(colors.HexColor("#F6F7F9"))
        canvas.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        logo_w = 4.7 * cm
        with Image.open(LOGO) as logo:
            logo_h = logo_w * (logo.height / logo.width)
        canvas.drawImage(
            str(LOGO),
            page_w - 1.7 * cm - logo_w,
            page_h - 1.35 * cm - logo_h,
            width=logo_w,
            height=logo_h,
            mask="auto",
        )
        canvas.setStrokeColor(colors.HexColor("#C91722"))
        canvas.setLineWidth(3)
        canvas.line(1.7 * cm, page_h - 4.25 * cm, page_w - 1.7 * cm, page_h - 4.25 * cm)
        canvas.setStrokeColor(colors.HexColor("#363A40"))
        canvas.setLineWidth(0.75)
        canvas.rect(1.35 * cm, 1.25 * cm, page_w - 2.7 * cm, page_h - 2.5 * cm, fill=0, stroke=1)
        canvas.setFont(PDF_FONT, 8)
        canvas.setFillColor(colors.HexColor("#6F747C"))
        canvas.drawCentredString(page_w / 2, 0.82 * cm, "Auto Sosnovec s.r.o. | Přístavní 27, Praha 7 | autososnovec@seznam.cz")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=1.7 * cm,
        leftMargin=1.7 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.55 * cm,
    )

    flow = []
    flow.append(Paragraph(f"<font name='{PDF_FONT_BOLD}' color='#6F747C'>Auto Sosnovec s.r.o.</font>", body))
    flow.append(Paragraph(f"<font name='{PDF_FONT_BOLD}' size='28'>Ceník služeb</font>", body))
    flow.append(Spacer(1, 30))

    flow.append(Paragraph("Mechanické, karosářské a elektrikářské práce", section_title))
    flow.append(make_pdf_table([["Služba", "Cena"], *ITEMS_BASIC], [11.1 * cm, 5.4 * cm]))

    flow.append(Paragraph("Diagnostika", section_title))
    diag_rows = [["Typ", "Rozsah práce", "Cena"]]
    for name, desc, price in ITEMS_DIAGNOSTICS:
        diag_rows.append([name, Paragraph(desc, body), price])
    flow.append(make_pdf_table(diag_rows, [4.4 * cm, 8.5 * cm, 3.6 * cm], 3))

    flow.append(Paragraph("Plnění klimatizací", section_title))
    flow.append(make_pdf_table([["Služba", "Cena"], *ITEMS_AC], [11.1 * cm, 5.4 * cm]))

    flow.append(Spacer(1, 14))
    flow.append(Paragraph(f"<font name='{PDF_FONT_BOLD}' color='#6F747C'>Ceny položek R134A a R1234yf jsou ponechány k ručnímu doplnění ve Word kopii.</font>", body))

    doc.build(flow, onFirstPage=draw_background, onLaterPages=draw_background)


if __name__ == "__main__":
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)
